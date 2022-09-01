import re
import json
import requests as requests
from docx import Document
import os.path
import textract


directory = "liste_personnes"
if not os.path.exists(directory):
    os.makedirs(directory)



#firstly, fetch data using the PuppetPlays API
url = 'https://api.puppetplays.eu/graphql/'
query_fr = """{ entries(typeId:5, relatedToEntries: {section: "works"}, site: "fr") { id slug title ... on persons_persons_Entry { firstName lastName nickname usualName birthDate deathDate gender biographicalNote } } }""" #{ ... on works_works_Entry authors { title } }
r_fr = requests.post(url, json={'query': query_fr})
query_en = """{ entries(typeId:5, site: "en") { id slug title ... on persons_persons_Entry { firstName lastName nickname usualName birthDate deathDate gender biographicalNote } } }""" #{ ... on works_works_Entry authors { title } }
r_en = requests.post(url, json={'query': query_en})

#load data and transform it
json_data = json.loads(r_fr.text)
liste_personnes = json_data['data']['entries']
json_data_en = json.loads(r_en.text)
liste_personnes_en = json_data_en['data']['entries']

liste_english_slug = []

for personne_en in liste_personnes_en:
    liste_english_slug.append(personne_en['slug'])

Number_of_words = 0
for personne in liste_personnes:
    if personne['slug'] not in liste_english_slug and personne["biographicalNote"] != None:
        #print(type(oeuvre), oeuvre)
        document = Document() #init word document to write


        #Écriture du titre de la page courante
        document.add_heading(personne['title'], 0)

        document.add_heading('Informations générales', level=1)

        #Retrait des balises pour une lecture humaine.
        pattern = '<[^>]*>'
        replace = ''
        resume = re.sub(pattern, replace, personne['biographicalNote'])

        #Ecris la notice.
        #Prénom
        p = document.add_paragraph('\n')
        p.add_run("Prénom : ").bold = True
        p.add_run(personne["firstName"]).bold = True
        #Nom
        p = document.add_paragraph('\n')
        p.add_run("Nom : ").bold = True
        p.add_run(personne["lastName"]).bold = True
        #Surnom
        p = document.add_paragraph('\n')
        p.add_run("Surnom : ").bold = True
        p.add_run(personne["nickname"]).bold = True
        # Nom d'usage
        p = document.add_paragraph('\n')
        p.add_run("Nom d'usage : ").bold = True
        p.add_run(personne["usualName"]).bold = True
        # Date de naissance
        p = document.add_paragraph('\n')
        p.add_run("Date de naissance : ").bold = True
        p.add_run(personne["birthDate"]).bold = True
        # Date de décès
        p = document.add_paragraph('\n')
        p.add_run("Date de décès : ").bold = True
        p.add_run(personne["deathDate"]).bold = True
        # Genre
        p = document.add_paragraph('\n')
        p.add_run("Genre : ").bold = True
        p.add_run(personne["gender"]).bold = True

        #Écriture du résumé.
        document.add_heading('Notice biographique', level=1)
        p = document.add_paragraph('\n')
        p.add_run(resume)

        #Ajout du titre et suppression des éléments causant des erreurs dans les noms de fichiers
        titre = personne['title']
        titre = titre.replace(" ", "_")
        titre = titre.replace("\'", "_")
        titre = titre.replace("?", "")
        titre = titre.replace(",", "")
        titre = titre.replace("-", "_")

        #initialise le nom du fichier
        creation = ".\\" + directory + "\\" + titre + ".docx"

        # permet de ne pas écraser les oeuvres ayant le même titre en leur ajoutant un numéro
        counter = 1

        #Compte le nombre de mots
        Number = len(resume.split())
        Number_of_words += Number
        while os.path.exists(creation):
            if resume in textract.process(creation).decode('utf-8'):
                break
            print(titre)
            counter += 1
            creation = ".\\" + directory + "\\" + titre + "_" + str(counter) + ".docx"
            print(counter)
        document.save(creation)
    else:
        print("Nous avons trouvé cette personne : ", personne['title'])

print("Il y a un total de : ", Number_of_words, " mots à traduire au total.")